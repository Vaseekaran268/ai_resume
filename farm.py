import streamlit as st
import fitz
import requests
from bs4 import BeautifulSoup
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.vectorstores import FAISS
from langchain.embeddings import FastEmbedEmbeddings
from langchain.memory import ConversationBufferMemory
from langchain.chains import ConversationalRetrievalChain
from langchain_groq import ChatGroq
import os
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime

# Function to extract text from a PDF file
def load_pdf_text(pdf_path):
    doc = fitz.open(pdf_path)
    text_data = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text_data.append(page.get_text("text"))

    return text_data

# Function to fetch and extract text from a URL
def fetch_url_content(url):
    try:
        response = requests.get(url)
        if response.status_code == 200:
            # Check if the URL points to a PDF
            if url.lower().endswith(".pdf"):
                with open("temp.pdf", "wb") as f:
                    f.write(response.content)
                return load_pdf_text("temp.pdf")
            else:
                # Extract text from HTML
                soup = BeautifulSoup(response.text, "html.parser")
                return [soup.get_text()]
        else:
            st.error(f"Failed to fetch URL: {response.status_code}")
            return []
    except Exception as e:
        st.error(f"Error fetching URL: {e}")
        return []

# Function to create a table with boxes for each character
def create_boxed_table(doc, label, value):
    table = doc.add_table(rows=1, cols=len(value))
    table.autofit = False
    table.allow_autofit = False
    for i, char in enumerate(value):
        cell = table.cell(0, i)
        cell.text = char
        cell.width = Pt(20)
        cell.paragraphs[0].alignment = 1  # Center alignment
        cell.vertical_alignment = 1  # Center alignment
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tcBorders.append(border)
        tcPr.append(tcBorders)
    doc.add_paragraph(label, style='Intense Quote')

# Function to generate a Word document for the Agri Loan Application
def generate_agri_loan_application(personal_info, land_details, loan_details, bank_details, photo, signature):
    doc = Document()
    doc.add_heading('Rythu-Mitra Agricultural Loan Application for Farmers', 0)

    # Add photo to the top-right side
    if photo is not None:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(photo, width=Inches(1.5))
        paragraph.alignment = 2  # Right alignment

    # Add Personal Information
    doc.add_heading('Personal Information', level=1)
    create_boxed_table(doc, "Full Name:", personal_info['name'])
    create_boxed_table(doc, "Date of Birth:", str(personal_info['dob']))
    create_boxed_table(doc, "Identity Proof Type:", personal_info['identity_proof_type'])
    create_boxed_table(doc, "Identity Proof Number:", personal_info['identity_proof'])
    create_boxed_table(doc, "Address Proof:", personal_info['address_proof'])
    create_boxed_table(doc, "Mobile Number:", personal_info['mobile'])
    create_boxed_table(doc, "Email ID:", personal_info['email'])

    # Add Land & Farming Details
    doc.add_heading('Land & Farming Details', level=1)
    create_boxed_table(doc, "Land Area:", land_details['land_area'])
    create_boxed_table(doc, "Location:", land_details['location'])
    create_boxed_table(doc, "Crop Type:", land_details['crop_type'])
    create_boxed_table(doc, "Past Yield:", land_details['past_yield'])
    create_boxed_table(doc, "Water Source:", land_details['water_source'])

    # Add Loan-Related Details
    doc.add_heading('Loan-Related Details', level=1)
    create_boxed_table(doc, "Loan Amount Required:", loan_details['loan_amount'])
    create_boxed_table(doc, "Purpose of Loan:", loan_details['purpose'])
    create_boxed_table(doc, "Repayment Period:", loan_details['repayment_period'])
    create_boxed_table(doc, "Previous Loan Details:", loan_details['previous_loan'])

    # Add Bank Details
    doc.add_heading('Bank Details', level=1)
    create_boxed_table(doc, "Bank Account Number:", bank_details['account_number'])

    # Add Declaration
    doc.add_heading('Declaration', level=1)
    doc.add_paragraph("I certify that the information given above and in the enclosures are true in all respects and that this shall form the basis of loan application.")

    # Add Signature to the bottom-right side
    if signature is not None:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(signature, width=Inches(1.5))
        paragraph.alignment = 2  # Right alignment

    # Add Name, Signature, and Date
    doc.add_paragraph(f"Name and Signature/Thumb Impression of the Applicant: {personal_info['name']}")
    doc.add_paragraph(f"Place: {personal_info['city']}")
    doc.add_paragraph(f"Date: {datetime.today().strftime('%Y-%m-%d')}")

    # Save the document to a BytesIO object
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    return doc_bytes

# Streamlit UI setup
st.set_page_config(page_title="PDF/URL Q&A Bot", layout="wide")
st.markdown(
    """
    <style>
    body {
        background-color: #008000;
        color: #333333;
    }
    .stButton > button {
        background-color: #4CAF50;
        color: white;
        border: none;
        padding: 10px 20px;
        font-size: 16px;
        cursor: pointer;
        border-radius: 5px;
    }
    .stButton > button:hover {
        background-color: #45a049;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

st.title("📄 PDF/URL Q&A Bot with Text Retrieval")

# Option to upload a PDF, submit a URL, or fill a loan application
option = st.radio("Choose an option:", ("Upload PDF", "Submit URL", "Agri Loan Application"))

if option == "Upload PDF":
    uploaded_file = st.file_uploader("Upload your PDF file", type=["pdf"])

    if uploaded_file is not None:
        with st.spinner("Processing your PDF..."):
            with open("uploaded_file.pdf", "wb") as f:
                f.write(uploaded_file.read())

            # Load text from the PDF
            texts = load_pdf_text("uploaded_file.pdf")

            # Prepare text chunks for embedding
            text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
            text_chunks = text_splitter.create_documents(texts)
            text_embeddings = FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")
            text_db = FAISS.from_documents(text_chunks, text_embeddings)

            # Setup language model for Q&A
            os.environ['GROQ_API_KEY'] = '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'
            llm = ChatGroq(model_name='llama-3.1-8b-instant')
            memory = ConversationBufferMemory(memory_key='chat_history', return_messages=False)
            retriever = text_db.as_retriever(search_type="similarity", search_kwargs={"k": 3})
            qa = ConversationalRetrievalChain.from_llm(
                llm=llm,
                memory=memory,
                retriever=retriever
            )

            st.success("PDF processed successfully!")

            # Sidebar for extracted text
            st.sidebar.subheader("📜 Extracted Text")
            for i, text in enumerate(texts):
                with st.sidebar.expander(f"Page {i + 1}"):
                    st.write(text)

            # Tabs for asking questions and exploring text
            tab1, tab2 = st.tabs(["💬 Ask Questions", "📜 Explore Text"])

            with tab1:
                query = st.text_input("Ask a question about the PDF:")
                if query:
                    with st.spinner("Searching for answers..."):
                        result = qa({"question": query})
                        answer = result['answer']
                        st.write("### 💡 Answer:", answer)

            with tab2:
                st.subheader("📜 Extracted Text")
                for i, text in enumerate(texts):
                    with st.expander(f"Page {i + 1}"):
                        st.write(text)

    else:
        st.info("Upload a PDF file to get started!")

elif option == "Submit URL":  # Submit URL
    url = st.text_input("Enter the URL:")

    if url:
        with st.spinner("Fetching and processing URL content..."):
            # Fetch and extract text from the URL
            texts = fetch_url_content(url)

            if texts:
                # Prepare text chunks for embedding
                text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=100)
                text_chunks = text_splitter.create_documents(texts)
                text_embeddings = FastEmbedEmbeddings(model_name="BAAI/bge-small-en-v1.5")
                text_db = FAISS.from_documents(text_chunks, text_embeddings)

                # Setup language model for Q&A
                os.environ['GROQ_API_KEY'] = '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'
                llm = ChatGroq(model_name='llama-3.1-8b-instant')
                memory = ConversationBufferMemory(memory_key='chat_history', return_messages=False)
                retriever = text_db.as_retriever(search_type="similarity", search_kwargs={"k": 3})
                qa = ConversationalRetrievalChain.from_llm(
                    llm=llm,
                    memory=memory,
                    retriever=retriever
                )

                st.success("URL content processed successfully!")

                # Sidebar for extracted text
                st.sidebar.subheader("📜 Extracted Text")
                for i, text in enumerate(texts):
                    with st.sidebar.expander(f"Section {i + 1}"):
                        st.write(text)

                # Tabs for asking questions and exploring text
                tab1, tab2 = st.tabs(["💬 Ask Questions", "📜 Explore Text"])

                with tab1:
                    query = st.text_input("Ask a question about the URL content:")
                    if query:
                        with st.spinner("Searching for answers..."):
                            result = qa({"question": query})
                            answer = result['answer']
                            st.write("### 💡 Answer:", answer)

                with tab2:
                    st.subheader("📜 Extracted Text")
                    for i, text in enumerate(texts):
                        with st.expander(f"Section {i + 1}"):
                            st.write(text)
            else:
                st.error("No text could be extracted from the URL.")

else:  # Agri Loan Application
    st.subheader("Agri Loan Application Form")

    # Personal Information
    st.header("Personal Information")
    name = st.text_input("Full Name:")
    dob = st.date_input("Date of Birth:")
    identity_proof_type = st.selectbox("Identity Proof Type:", ["Aadhaar Card", "Voter ID", "PAN Card"])
    identity_proof = st.text_input(f"{identity_proof_type} Number:")
    address_proof = st.text_input("Address Proof (Ration Card, Utility Bill, etc.):")
    mobile = st.text_input("Mobile Number:")
    email = st.text_input("Email ID:")
    photo = st.file_uploader("Upload Passport-size Photograph", type=["jpg", "jpeg", "png"])
    signature = st.file_uploader("Upload Signature/Thumb Impression", type=["jpg", "jpeg", "png"])

    # Land & Farming Details
    st.header("Land & Farming Details")
    land_area = st.text_input("Land Area:")
    location = st.text_input("Location:")
    crop_type = st.text_input("Type of Crop Grown:")
    past_yield = st.text_input("Past Yield (if any):")
    water_source = st.selectbox("Water Source:", ["Borewell", "Canal", "Rain-fed"])

    # Loan-Related Details
    st.header("Loan-Related Details")
    loan_amount = st.text_input("Loan Amount Required:")
    purpose = st.selectbox("Purpose of Loan:", ["Crop Cultivation", "Equipment Purchase", "Irrigation", "Dairy", "Poultry"])
    repayment_period = st.text_input("Repayment Period:")
    previous_loan = st.text_input("Previous Loan Details (if any):")

    # Bank Details
    st.header("Bank Details")
    account_number = st.text_input("Bank Account Number:")

    if st.button("Submit Application"):
        if (name and dob and identity_proof and address_proof and mobile and email and
            land_area and location and crop_type and water_source and
            loan_amount and purpose and repayment_period and account_number):
            # Prepare data for document generation
            personal_info = {
                "name": name,
                "dob": dob,
                "identity_proof_type": identity_proof_type,
                "identity_proof": identity_proof,
                "address_proof": address_proof,
                "mobile": mobile,
                "email": email,
                "city": location.split(",")[-1].strip() if location else ""
            }
            land_details = {
                "land_area": land_area,
                "location": location,
                "crop_type": crop_type,
                "past_yield": past_yield,
                "water_source": water_source
            }
            loan_details = {
                "loan_amount": loan_amount,
                "purpose": purpose,
                "repayment_period": repayment_period,
                "previous_loan": previous_loan
            }
            bank_details = {
                "account_number": account_number
            }

            # Generate the Agri Loan Application document
            doc_bytes = generate_agri_loan_application(personal_info, land_details, loan_details, bank_details, photo, signature)

            # Provide download link
            st.success("Application submitted successfully!")
            st.download_button(
                label="Download Application Form",
                data=doc_bytes,
                file_name="agri_loan_application.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Please fill in all the required fields.")